# -*- coding: utf-8 -*-
"""
ML 보고서 Word 문서 생성 (한글 폰트 지원)
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import pandas as pd
from pathlib import Path
from datetime import datetime

# 경로
BASE_DIR = Path(__file__).parent
ML_DATA_DIR = Path(r"C:\Users\hsh\Desktop\vibecoding\주식웹사이트\국내주식웹사이트\stock-screener-kr\ml\data")


def set_cell_shading(cell, color):
    """셀 배경색 설정"""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_korean_font(run, font_name='맑은 고딕', size=10):
    """한글 폰트 설정"""
    run.font.name = font_name
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)


def add_paragraph_korean(doc, text, bold=False, size=10):
    """한글 문단 추가"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    set_korean_font(run, size=size)
    return p


def add_heading_korean(doc, text, level=1):
    """한글 제목 추가"""
    heading = doc.add_heading(level=level)
    run = heading.add_run(text)
    set_korean_font(run, size=14 if level == 1 else 12)
    return heading


def add_table_with_header(doc, headers, data, header_color="4472C4"):
    """헤더가 있는 테이블 추가"""
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_korean_font(run, size=9)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)

    # 데이터
    for row_data in data:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(str(value))
            set_korean_font(run, size=9)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    return table


def create_report():
    doc = Document()

    # 기본 폰트 설정
    style = doc.styles['Normal']
    style.font.name = '맑은 고딕'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

    # ========== 제목 ==========
    title = doc.add_heading(level=0)
    run = title.add_run('52주 신고가 돌파 종목 ML 예측 모델 개발 보고서')
    set_korean_font(run, size=18)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 작성 정보
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run1 = p.add_run(f'작성일: {datetime.now().strftime("%Y년 %m월 %d일")}\n')
    set_korean_font(run1, size=10)
    run2 = p.add_run('버전: 1.0')
    set_korean_font(run2, size=10)

    doc.add_paragraph()

    # ========== 1. 프로젝트 개요 ==========
    add_heading_korean(doc, '1. 프로젝트 개요', level=1)

    add_heading_korean(doc, '1.1 목적', level=2)
    add_paragraph_korean(doc, '52주 신고가를 돌파한 종목에 대해 머신러닝 모델을 활용하여 다음을 예측합니다:')
    add_paragraph_korean(doc, '  • 예상 최고 수익률: 돌파 후 20거래일 내 도달할 최고 수익률')
    add_paragraph_korean(doc, '  • 성공 확률: 10% 이상 상승할 확률')

    add_heading_korean(doc, '1.2 배경', level=2)
    p = doc.add_paragraph()
    run1 = p.add_run('52주 신고가 돌파')
    run1.bold = True
    set_korean_font(run1)
    run2 = p.add_run('는 기술적 분석에서 강력한 매수 신호로 알려져 있습니다. 그러나 모든 돌파가 성공하지는 않으며, 돌파의 "품질"에 따라 결과가 달라집니다. 본 프로젝트는 과거 3년간의 돌파 데이터를 학습하여 돌파 시점의 특성만으로 향후 수익률을 예측하는 모델을 개발하였습니다.')
    set_korean_font(run2)

    add_heading_korean(doc, '1.3 시스템 구성도', level=2)
    add_paragraph_korean(doc, '[데이터 수집] → [피처 엔지니어링] → [모델 학습] → [실시간 예측]', bold=True)
    add_paragraph_korean(doc, '   3년간 돌파        14개 기술적         LightGBM        웹사이트')
    add_paragraph_korean(doc, '  이벤트 수집         지표 계산         회귀+분류         표시')

    # ========== 2. 데이터 수집 ==========
    add_heading_korean(doc, '2. 데이터 수집', level=1)

    add_heading_korean(doc, '2.1 수집 기간 및 범위', level=2)
    headers = ['항목', '내용']
    data = [
        ['수집 기간', '2022년 1월 ~ 2025년 6월 (약 3.5년)'],
        ['대상 시장', 'KOSPI, KOSDAQ'],
        ['시가총액 필터', '1,000억원 이상'],
        ['제외 종목', 'ETF, 스팩, 리츠 등'],
    ]
    add_table_with_header(doc, headers, data)
    doc.add_paragraph()

    add_heading_korean(doc, '2.2 돌파 및 라벨 정의', level=2)
    p = doc.add_paragraph()
    run1 = p.add_run('돌파 조건: ')
    run1.bold = True
    set_korean_font(run1)
    run2 = p.add_run('당일 종가 > 직전 250거래일(52주) 최고가')
    set_korean_font(run2)

    p = doc.add_paragraph()
    run1 = p.add_run('성공 라벨 (label=1): ')
    run1.bold = True
    set_korean_font(run1)
    run2 = p.add_run('돌파일 종가 대비 20거래일 내 최고가가 10% 이상 상승')
    set_korean_font(run2)

    add_heading_korean(doc, '2.3 수집된 데이터 요약', level=2)
    headers = ['항목', '값']
    data = [
        ['총 샘플 수', '4,042건'],
        ['성공 비율 (Baseline)', '40.4%'],
        ['평균 최고 수익률', '14.43%'],
    ]
    add_table_with_header(doc, headers, data)
    doc.add_paragraph()

    # ========== 3. 피처 엔지니어링 ==========
    add_heading_korean(doc, '3. 피처 엔지니어링', level=1)

    add_heading_korean(doc, '3.1 피처 설계 철학', level=2)
    add_paragraph_korean(doc, '돌파의 "품질"을 다각도로 평가할 수 있는 14개 피처를 6개 카테고리로 설계하였습니다.')

    add_heading_korean(doc, '3.2 피처 목록 (14개)', level=2)

    # A. 돌파 품질
    add_paragraph_korean(doc, 'A. 돌파 품질 (3개)', bold=True)
    headers = ['피처명', '설명', '계산 방법']
    data = [
        ['breakout_pct', '돌파 강도', '(돌파일 종가 - 52주 고가) / 52주 고가 × 100'],
        ['volume_surge', '거래량 급증', '돌파일 거래량 / 50일 평균 거래량'],
        ['close_strength', '종가 강도', '(종가 - 저가) / (고가 - 저가)'],
    ]
    add_table_with_header(doc, headers, data, "2E75B6")
    doc.add_paragraph()

    # B. 베이스 품질
    add_paragraph_korean(doc, 'B. 베이스 품질 (3개)', bold=True)
    headers = ['피처명', '설명', '계산 방법']
    data = [
        ['base_length', '베이스 기간', '직전 52주 고가 이후 경과일'],
        ['volatility_contraction', '변동성 수축', '최근 5일 변동폭 / 20일 변동폭'],
        ['volume_dry_up', '거래량 고갈', '최근 10일 평균 / 50일 평균 거래량'],
    ]
    add_table_with_header(doc, headers, data, "2E75B6")
    doc.add_paragraph()

    # C. 추세 상태
    add_paragraph_korean(doc, 'C. 추세 상태 (2개)', bold=True)
    headers = ['피처명', '설명', '계산 방법']
    data = [
        ['ma200_slope', '장기 추세', '(현재 200MA - 20일전 200MA) / 20일전 200MA'],
        ['pct_above_52w_low', '저점 대비 상승률', '(종가 - 52주 저가) / 52주 저가 × 100'],
    ]
    add_table_with_header(doc, headers, data, "2E75B6")
    doc.add_paragraph()

    # D. 상대강도
    add_paragraph_korean(doc, 'D. 상대강도 (2개)', bold=True)
    headers = ['피처명', '설명', '계산 방법']
    data = [
        ['rs_vs_market', '시장 대비 강도', '종목 20일 수익률 - 시장 20일 수익률'],
        ['market_return', '시장 상황', 'KOSPI/KOSDAQ 20일 수익률'],
    ]
    add_table_with_header(doc, headers, data, "2E75B6")
    doc.add_paragraph()

    # E. 리스크
    add_paragraph_korean(doc, 'E. 리스크 지표 (2개)', bold=True)
    headers = ['피처명', '설명', '계산 방법']
    data = [
        ['ma20_deviation', '이격도', '종가 / 20일 이동평균 - 1'],
        ['liquidity', '유동성', '5일 평균 거래대금'],
    ]
    add_table_with_header(doc, headers, data, "2E75B6")
    doc.add_paragraph()

    # F. 저항선/변동성
    add_paragraph_korean(doc, 'F. 저항선/변동성 (2개)', bold=True)
    headers = ['피처명', '설명', '계산 방법']
    data = [
        ['days_since_ath', 'ATH 이후 경과일', '역사적 신고가 이후 경과일'],
        ['atr_ratio', 'ATR 비율', 'ATR(14) / 현재가 × 100'],
    ]
    add_table_with_header(doc, headers, data, "2E75B6")
    doc.add_paragraph()

    # ========== 4. 모델 비교 ==========
    add_heading_korean(doc, '4. 모델 선정', level=1)

    add_heading_korean(doc, '4.1 비교 대상', level=2)
    add_paragraph_korean(doc, '7개 모델을 3가지 피처셋으로 비교하였습니다.')

    headers = ['모델 유형', '모델명']
    data = [
        ['선형 모델', 'Ridge, Lasso, ElasticNet'],
        ['앙상블 (배깅)', 'Random Forest'],
        ['앙상블 (부스팅)', 'XGBoost, LightGBM, Gradient Boosting'],
        ['앙상블 (스태킹)', 'Voting, Stacking'],
    ]
    add_table_with_header(doc, headers, data)
    doc.add_paragraph()

    add_heading_korean(doc, '4.2 회귀 모델 성능 비교 (R² 기준 상위 10개)', level=2)

    # 실제 데이터 로드
    try:
        df_reg = pd.read_csv(ML_DATA_DIR / "model_comparison_results.csv", encoding='utf-8-sig')
        df_reg = df_reg.sort_values('R2', ascending=False).head(10)

        headers = ['피처셋', '모델', 'MAE', 'RMSE', 'R²']
        data = []
        for _, row in df_reg.iterrows():
            data.append([
                row['feature_set'],
                row['model'],
                f"{row['MAE']:.2f}",
                f"{row['RMSE']:.2f}",
                f"{row['R2']:.4f}"
            ])
        add_table_with_header(doc, headers, data, "70AD47")
    except Exception as e:
        add_paragraph_korean(doc, f"(데이터 로드 실패: {e})")
    doc.add_paragraph()

    # 회귀 모델 성능 해석
    p = doc.add_paragraph()
    run1 = p.add_run('해석: ')
    run1.bold = True
    set_korean_font(run1)
    run2 = p.add_run('Lasso와 Ridge 모델이 R² 기준으로 가장 좋은 성능을 보였습니다. 그러나 튜닝 후 LightGBM이 R² 0.0343으로 가장 높은 성능을 달성하여 최종 모델로 선정되었습니다.')
    set_korean_font(run2)

    add_heading_korean(doc, '4.3 분류 모델 성능 비교 (Precision 기준)', level=2)

    try:
        df_cls = pd.read_csv(ML_DATA_DIR / "classification_results.csv", encoding='utf-8-sig')
        df_cls = df_cls.sort_values('Precision', ascending=False)

        headers = ['피처셋', '모델', 'Accuracy', 'Precision', 'Recall', 'F1', 'AUC']
        data = []
        for _, row in df_cls.iterrows():
            data.append([
                row['feature_set'],
                row['model'],
                f"{row['Accuracy']:.4f}",
                f"{row['Precision']:.4f}",
                f"{row['Recall']:.4f}",
                f"{row['F1']:.4f}",
                f"{row['AUC']:.4f}"
            ])
        add_table_with_header(doc, headers, data, "C55A11")
    except Exception as e:
        add_paragraph_korean(doc, f"(데이터 로드 실패: {e})")
    doc.add_paragraph()

    # 분류 모델 성능 해석
    p = doc.add_paragraph()
    run1 = p.add_run('해석: ')
    run1.bold = True
    set_korean_font(run1)
    run2 = p.add_run('LogisticRegression이 Precision 55.6%로 가장 높지만 Recall이 17.6%로 매우 낮습니다. LightGBM은 Precision 50.0%, Recall 32.9%로 균형 잡힌 성능을 보여 최종 모델로 선정되었습니다.')
    set_korean_font(run2)

    add_heading_korean(doc, '4.4 하이퍼파라미터 튜닝 결과', level=2)

    try:
        df_tune = pd.read_csv(ML_DATA_DIR / "tuning_results.csv", encoding='utf-8-sig')
        df_tune = df_tune.sort_values('R2', ascending=False).head(5)

        headers = ['피처셋', '모델', 'MAE', 'R²', 'Precision', 'AUC']
        data = []
        for _, row in df_tune.iterrows():
            data.append([
                row['feature_set'],
                row['model'],
                f"{row['MAE']:.2f}",
                f"{row['R2']:.4f}",
                f"{row['Precision']:.4f}",
                f"{row['AUC']:.4f}"
            ])
        add_table_with_header(doc, headers, data, "7030A0")
    except Exception as e:
        add_paragraph_korean(doc, f"(데이터 로드 실패: {e})")
    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('최종 선정: ')
    run1.bold = True
    set_korean_font(run1)
    run2 = p.add_run('튜닝된 LightGBM (피처셋 A, 14개 피처)')
    set_korean_font(run2)

    # ========== 5. 최종 모델 ==========
    add_heading_korean(doc, '5. 최종 모델 성능', level=1)

    add_heading_korean(doc, '5.1 회귀 모델 (예상 최고 수익률 예측)', level=2)

    headers = ['지표', '값', '해석']
    data = [
        ['MAE', '14.60%', '평균 예측 오차 ±14.60%p'],
        ['RMSE', '27.83%', '큰 오차에 민감한 지표'],
        ['R²', '0.0293 (2.9%)', '전체 변동의 2.9% 설명'],
    ]
    add_table_with_header(doc, headers, data, "4472C4")
    doc.add_paragraph()

    add_heading_korean(doc, '5.2 분류 모델 (성공 확률 예측)', level=2)

    headers = ['지표', '값', '해석']
    data = [
        ['Accuracy', '59.6%', '전체 정확도'],
        ['Precision', '50.1%', '예측이 맞을 확률 (핵심 지표)'],
        ['Recall', '32.9%', '실제 성공 중 포착률'],
        ['F1', '39.7%', '조화 평균'],
        ['AUC', '0.607', '분류 능력'],
    ]
    add_table_with_header(doc, headers, data, "4472C4")
    doc.add_paragraph()

    add_heading_korean(doc, '5.3 Baseline 대비 개선', level=2)

    headers = ['항목', 'Baseline', '모델 예측', '개선']
    data = [
        ['성공 확률', '40.4%', '50.1%', '+9.7%p (24% 향상)'],
    ]
    add_table_with_header(doc, headers, data, "00B050")
    doc.add_paragraph()

    p = doc.add_paragraph()
    run1 = p.add_run('핵심 인사이트: ')
    run1.bold = True
    set_korean_font(run1)
    run2 = p.add_run('무작위로 종목을 선택하면 40.4%가 성공하지만, 모델이 "성공"으로 예측한 종목만 선택하면 50.1%가 성공합니다. 이는 약 24%의 성공률 향상 효과입니다.')
    set_korean_font(run2)

    # ========== 6. 모델 해석 ==========
    add_heading_korean(doc, '6. 모델 해석', level=1)

    add_heading_korean(doc, '6.1 R² = 2.9%의 의미', level=2)
    add_paragraph_korean(doc, '  • 주식 수익률 예측은 본질적으로 매우 어려운 문제입니다.')
    add_paragraph_korean(doc, '  • 효율적 시장 가설(EMH)에 따르면 과거 정보로 미래 수익률 예측이 불가능합니다.')
    add_paragraph_korean(doc, '  • 2.9%의 설명력은 낮아 보이지만, 금융 데이터에서는 의미 있는 수준입니다.')
    add_paragraph_korean(doc, '  • 대부분의 퀀트 전략도 유사한 수준의 예측력을 가집니다.')

    add_heading_korean(doc, '6.2 모델의 한계', level=2)
    add_paragraph_korean(doc, '  1. 낮은 설명력: R² 2.9%는 대부분의 변동을 설명하지 못함')
    add_paragraph_korean(doc, '  2. 노이즈가 많은 데이터: 주가는 뉴스, 시장 심리 등 예측 불가 요인에 영향')
    add_paragraph_korean(doc, '  3. 과거 ≠ 미래: 과거 패턴이 미래에도 유효하다는 보장 없음')
    add_paragraph_korean(doc, '  4. 시장 환경 변화: 상승장/하락장에 따라 모델 성능 차이')

    add_heading_korean(doc, '6.3 모델의 활용 가치', level=2)
    add_paragraph_korean(doc, '  1. 종목 선별 도구: 확률이 높은 종목에 집중 가능')
    add_paragraph_korean(doc, '  2. 리스크 관리: 낮은 확률 종목 회피')
    add_paragraph_korean(doc, '  3. 정량적 기준: 주관적 판단 대신 객관적 지표 제공')

    # ========== 7. 서비스 적용 ==========
    add_heading_korean(doc, '7. 서비스 적용', level=1)

    add_heading_korean(doc, '7.1 추천 로직', level=2)

    headers = ['조건', '추천']
    data = [
        ['성공확률 ≥ 60% AND 예상수익률 ≥ 10%', 'BUY'],
        ['성공확률 ≥ 50% OR 예상수익률 ≥ 5%', 'HOLD'],
        ['그 외', 'PASS'],
    ]
    add_table_with_header(doc, headers, data)
    doc.add_paragraph()

    add_heading_korean(doc, '7.2 웹사이트 표시 예시', level=2)

    headers = ['종목명', '돌파일', '돌파강도', '거래량비', '예상 최고수익', '성공 확률', '추천']
    data = [
        ['현대시멘트', '2월2일', '+22.7%', '37.9x', '+38.6%', '87.0%', 'BUY'],
        ['삼양옵틱스', '1월30일', '+44.0%', '49.9x', '+42.7%', '87.8%', 'BUY'],
        ['웅진씽크빅', '2월2일', '+1.7%', '2.1x', '+21.7%', '79.0%', 'BUY'],
    ]
    add_table_with_header(doc, headers, data)
    doc.add_paragraph()

    # ========== 8. 결론 ==========
    add_heading_korean(doc, '8. 결론', level=1)

    add_heading_korean(doc, '8.1 주요 성과', level=2)
    add_paragraph_korean(doc, '  1. 14개 기술적 피처 설계 및 4,042건의 학습 데이터셋 구축')
    add_paragraph_korean(doc, '  2. LightGBM 기반 회귀/분류 모델 개발')
    add_paragraph_korean(doc, '  3. 종목 선별 효과 입증 (baseline 40.4% → 모델 예측 50.1%)')

    add_heading_korean(doc, '8.2 모델 요약', level=2)

    headers = ['항목', '회귀 모델', '분류 모델']
    data = [
        ['알고리즘', 'LightGBM Regressor', 'LightGBM Classifier'],
        ['타겟', '20일 내 최고 수익률', '10% 이상 상승 여부'],
        ['주요 지표', 'R² = 2.9%', 'Precision = 50.1%'],
        ['출력', 'predicted_gain (%)', 'success_probability (%)'],
    ]
    add_table_with_header(doc, headers, data)
    doc.add_paragraph()

    add_heading_korean(doc, '8.3 향후 개선 방향', level=2)
    add_paragraph_korean(doc, '  • 추가 피처: 섹터 정보, 외국인/기관 수급, 뉴스 센티먼트')
    add_paragraph_korean(doc, '  • 시계열 모델: LSTM, Transformer 활용')
    add_paragraph_korean(doc, '  • 앙상블 강화: 시장 환경별 다른 모델 적용')
    add_paragraph_korean(doc, '  • 백테스팅: 실제 투자 시뮬레이션으로 수익률 검증')

    # 저장
    output_path = BASE_DIR / "ML_예측모델_개발보고서.docx"
    doc.save(output_path)
    print(f"보고서 생성 완료: {output_path}")

    return output_path


if __name__ == "__main__":
    create_report()
